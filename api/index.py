import tempfile
import os
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import shutil
import uuid
from typing import Optional
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

app = FastAPI()

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Vercel serverless: use /tmp for temp files
UPLOAD_DIR = Path(tempfile.mkdtemp())
OUTPUT_DIR = Path(tempfile.mkdtemp())


@app.get("/")
async def root():
    return {
        "message": "文档转换API服务",
        "version": "1.0.0",
    }


@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "service": "document-converter"}


@app.post("/api/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    """Word转PDF"""
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="请上传.docx格式的Word文件")

        file_id = str(uuid.uuid4())
        docx_path = UPLOAD_DIR / f"{file_id}.docx"
        pdf_path = OUTPUT_DIR / f"{file_id}.pdf"

        with open(docx_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        doc = Document(str(docx_path))
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        width, height = letter
        margin = 50
        line_height = 14
        y_position = height - margin

        if doc.paragraphs and doc.paragraphs[0].text.strip():
            title = doc.paragraphs[0].text.strip()
            c.setFont("Helvetica-Bold", 16)
            c.drawString(margin, y_position, title)
            y_position -= line_height * 2
        else:
            c.setFont("Helvetica-Bold", 16)
            c.drawString(margin, y_position, "Document")
            y_position -= line_height * 2

        c.setFont("Helvetica", 11)
        for para in doc.paragraphs[1:]:
            text = para.text.strip()
            if not text:
                y_position -= line_height
                continue
            words = text.split()
            current_line = ""
            for word in words:
                test_line = current_line + word + " "
                if c.stringWidth(test_line, "Helvetica", 11) > (width - 2 * margin):
                    if y_position < margin:
                        c.showPage()
                        y_position = height - margin
                    c.drawString(margin, y_position, current_line.strip())
                    y_position -= line_height
                    current_line = word + " "
                else:
                    current_line = test_line
            if current_line.strip():
                if y_position < margin:
                    c.showPage()
                    y_position = height - margin
                c.drawString(margin, y_position, current_line.strip())
                y_position -= line_height

        c.save()

        try:
            docx_path.unlink()
        except:
            pass

        return FileResponse(
            path=str(pdf_path),
            media_type="application/pdf",
            filename=file.filename.replace(".docx", ".pdf")
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


@app.post("/api/convert/image-to-pdf")
async def image_to_pdf(file: UploadFile = File(...)):
    """图片转PDF"""
    try:
        file_id = str(uuid.uuid4())
        image_path = UPLOAD_DIR / f"{file_id}{Path(file.filename).suffix}"
        pdf_path = OUTPUT_DIR / f"{file_id}.pdf"

        with open(image_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        img = Image.open(image_path)
        img_width, img_height = img.size
        c = canvas.Canvas(str(pdf_path), pagesize=(img_width, img_height))
        c.drawImage(str(image_path), 0, 0, width=img_width, height=img_height)
        c.save()
        img.close()

        import time
        time.sleep(0.1)

        try:
            image_path.unlink()
        except:
            pass

        return FileResponse(
            path=str(pdf_path),
            media_type="application/pdf",
            filename=file.filename.replace(Path(file.filename).suffix, ".pdf")
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


@app.post("/api/convert/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    """PDF转Word"""
    raise HTTPException(status_code=501, detail="PDF转Word功能在云部署中暂不可用")


@app.post("/api/convert/pdf-to-image")
async def pdf_to_image(file: UploadFile = File(...)):
    """PDF转图片"""
    raise HTTPException(status_code=501, detail="PDF转图片功能在云部署中暂不可用")
