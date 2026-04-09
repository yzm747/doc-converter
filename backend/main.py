from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import shutil
from pathlib import Path
import uuid
from typing import Optional
import tempfile
import json
import hashlib
from datetime import datetime, timedelta
from pydantic import BaseModel

# 文档转换导入
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# PDF转换库（可选）
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    print("警告：pdf2docx未安装，PDF转Word功能不可用")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("警告：pdf2image未安装，PDF转图片功能不可用")

app = FastAPI(title="文档转换API", version="1.1.0")

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 创建临时目录
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
DATA_DIR = Path("data")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)

# ============ 用户数据存储 ============

USERS_FILE = DATA_DIR / "users.json"
ADMIN_PASSWORD = "admin888"  # 管理员密码，部署后请修改

def load_users():
    """加载用户数据"""
    if not USERS_FILE.exists():
        return {}
    try:
        with open(USERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {}

def save_users(users):
    """保存用户数据"""
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)

def hash_password(password):
    """密码哈希"""
    return hashlib.sha256(password.encode()).hexdigest()

# ============ 数据模型 ============

class RegisterRequest(BaseModel):
    phone: str
    email: str
    password: str

class LoginRequest(BaseModel):
    account: str  # 手机号或邮箱
    password: str

class AdminLoginRequest(BaseModel):
    password: str

class SetVipRequest(BaseModel):
    user_id: str
    days: int  # VIP天数

# ============ 用户接口 ============

@app.post("/api/user/register")
async def register(req: RegisterRequest):
    """用户注册"""
    users = load_users()

    # 验证手机号格式
    if req.phone and not req.phone.startswith("1") or (req.phone and len(req.phone) != 11):
        raise HTTPException(status_code=400, detail="手机号格式不正确")

    # 验证邮箱格式
    if req.email and "@" not in req.email:
        raise HTTPException(status_code=400, detail="邮箱格式不正确")

    # 检查手机号是否已注册
    if req.phone and req.phone in users:
        raise HTTPException(status_code=400, detail="该手机号已注册")

    # 检查邮箱是否已注册
    for uid, u in users.items():
        if u.get("email") == req.email:
            raise HTTPException(status_code=400, detail="该邮箱已注册")

    # 创建用户
    user_id = str(uuid.uuid4())[:8]
    today = datetime.now().strftime("%Y-%m-%d")

    users[user_id] = {
        "phone": req.phone or "",
        "email": req.email or "",
        "password": hash_password(req.password),
        "is_vip": False,
        "vip_expire": "",
        "created_at": today,
        "last_login": today,
        "daily_count": 0,
        "daily_date": today,
        "total_count": 0
    }

    save_users(users)

    return {
        "success": True,
        "user_id": user_id,
        "message": "注册成功"
    }

@app.post("/api/user/login")
async def login(req: LoginRequest):
    """用户登录"""
    users = load_users()

    # 查找用户
    found_user = None
    found_id = None
    for uid, u in users.items():
        if u.get("phone") == req.account or u.get("email") == req.account:
            found_user = u
            found_id = uid
            break

    if not found_user:
        raise HTTPException(status_code=400, detail="账号不存在")

    if found_user["password"] != hash_password(req.password):
        raise HTTPException(status_code=400, detail="密码错误")

    # 更新登录时间
    today = datetime.now().strftime("%Y-%m-%d")
    found_user["last_login"] = today
    save_users(users)

    # 检查VIP是否过期
    is_vip = found_user.get("is_vip", False)
    vip_expire = found_user.get("vip_expire", "")
    if is_vip and vip_expire and vip_expire < today:
        found_user["is_vip"] = False
        found_user["vip_expire"] = ""
        save_users(users)
        is_vip = False

    return {
        "success": True,
        "user_id": found_id,
        "phone": found_user.get("phone", ""),
        "email": found_user.get("email", ""),
        "is_vip": is_vip,
        "vip_expire": vip_expire if is_vip else "",
        "daily_count": found_user.get("daily_count", 0) if found_user.get("daily_date") == today else 0,
        "daily_limit": 3 if not is_vip else 999
    }

@app.get("/api/user/info/{user_id}")
async def get_user_info(user_id: str):
    """获取用户信息"""
    users = load_users()

    if user_id not in users:
        raise HTTPException(status_code=400, detail="用户不存在")

    u = users[user_id]
    today = datetime.now().strftime("%Y-%m-%d")

    # 检查VIP过期
    is_vip = u.get("is_vip", False)
    vip_expire = u.get("vip_expire", "")
    if is_vip and vip_expire and vip_expire < today:
        u["is_vip"] = False
        u["vip_expire"] = ""
        save_users(users)
        is_vip = False

    # 检查日期重置
    if u.get("daily_date") != today:
        u["daily_count"] = 0
        u["daily_date"] = today
        save_users(users)

    return {
        "user_id": user_id,
        "phone": u.get("phone", ""),
        "email": u.get("email", ""),
        "is_vip": is_vip,
        "vip_expire": vip_expire if is_vip else "",
        "daily_count": u.get("daily_count", 0),
        "daily_limit": 3 if not is_vip else 999,
        "total_count": u.get("total_count", 0)
    }

@app.post("/api/user/increment/{user_id}")
async def increment_count(user_id: str):
    """增加转换次数"""
    users = load_users()

    if user_id not in users:
        raise HTTPException(status_code=400, detail="用户不存在")

    u = users[user_id]
    today = datetime.now().strftime("%Y-%m-%d")

    # 检查VIP过期
    is_vip = u.get("is_vip", False)
    vip_expire = u.get("vip_expire", "")
    if is_vip and vip_expire and vip_expire < today:
        u["is_vip"] = False
        u["vip_expire"] = ""
        save_users(users)
        is_vip = False

    # 检查日期重置
    if u.get("daily_date") != today:
        u["daily_count"] = 0
        u["daily_date"] = today

    # 检查次数限制
    if not is_vip and u["daily_count"] >= 3:
        return {
            "success": False,
            "message": "今日免费次数已用完（3次），请联系客服开通VIP"
        }

    u["daily_count"] += 1
    u["total_count"] = u.get("total_count", 0) + 1
    save_users(users)

    remaining = 999 if is_vip else (3 - u["daily_count"])
    return {
        "success": True,
        "daily_count": u["daily_count"],
        "remaining": remaining,
        "is_vip": is_vip
    }

# ============ 管理员接口 ============

@app.post("/api/admin/login")
async def admin_login(req: AdminLoginRequest):
    """管理员登录"""
    if req.password != ADMIN_PASSWORD:
        raise HTTPException(status_code=400, detail="管理员密码错误")
    return {"success": True}

@app.get("/api/admin/users")
async def admin_get_users():
    """获取所有用户列表"""
    users = load_users()
    today = datetime.now().strftime("%Y-%m-%d")

    user_list = []
    for uid, u in users.items():
        is_vip = u.get("is_vip", False)
        vip_expire = u.get("vip_expire", "")

        # 自动过期检查
        if is_vip and vip_expire and vip_expire < today:
            u["is_vip"] = False
            u["vip_expire"] = ""
            is_vip = False

        user_list.append({
            "user_id": uid,
            "phone": u.get("phone", ""),
            "email": u.get("email", ""),
            "is_vip": is_vip,
            "vip_expire": vip_expire,
            "created_at": u.get("created_at", ""),
            "last_login": u.get("last_login", ""),
            "total_count": u.get("total_count", 0)
        })

    save_users(users)
    return {"users": user_list}

@app.post("/api/admin/set-vip")
async def admin_set_vip(req: SetVipRequest):
    """设置VIP（管理员操作）"""
    users = load_users()

    if req.user_id not in users:
        raise HTTPException(status_code=400, detail="用户不存在")

    u = users[req.user_id]
    today = datetime.now()

    # 如果已经是VIP，在原到期时间基础上续期
    if u.get("is_vip") and u.get("vip_expire"):
        try:
            current_expire = datetime.strptime(u["vip_expire"], "%Y-%m-%d")
            if current_expire > today:
                new_expire = current_expire + timedelta(days=req.days)
            else:
                new_expire = today + timedelta(days=req.days)
        except:
            new_expire = today + timedelta(days=req.days)
    else:
        new_expire = today + timedelta(days=req.days)

    u["is_vip"] = True
    u["vip_expire"] = new_expire.strftime("%Y-%m-%d")
    save_users(users)

    return {
        "success": True,
        "message": f"已为用户 {u.get('phone') or u.get('email')} 开通VIP {req.days}天",
        "vip_expire": u["vip_expire"]
    }

@app.post("/api/admin/remove-vip/{user_id}")
async def admin_remove_vip(user_id: str):
    """取消VIP"""
    users = load_users()

    if user_id not in users:
        raise HTTPException(status_code=400, detail="用户不存在")

    users[user_id]["is_vip"] = False
    users[user_id]["vip_expire"] = ""
    save_users(users)

    return {"success": True, "message": "已取消VIP"}

@app.delete("/api/admin/delete-user/{user_id}")
async def admin_delete_user(user_id: str):
    """删除用户"""
    users = load_users()

    if user_id not in users:
        raise HTTPException(status_code=400, detail="用户不存在")

    del users[user_id]
    save_users(users)

    return {"success": True, "message": "已删除用户"}

# ============ VIP套餐接口 ============

@app.get("/api/vip/plans")
async def get_vip_plans():
    """获取VIP套餐列表"""
    return {
        "plans": [
            {
                "id": "week",
                "name": "7天VIP",
                "price": 3.9,
                "original_price": 9.9,
                "days": 7,
                "tag": "超值",
                "features": ["无限次转换", "优先处理", "100MB文件限制"]
            },
            {
                "id": "month",
                "name": "月度VIP",
                "price": 9.9,
                "original_price": 29.9,
                "days": 30,
                "tag": "热门",
                "features": ["无限次转换", "优先处理", "200MB文件限制", "专属客服"]
            },
            {
                "id": "year",
                "name": "年度VIP",
                "price": 59.9,
                "original_price": 118.8,
                "days": 365,
                "tag": "推荐",
                "features": ["无限次转换", "最高优先", "500MB文件限制", "专属客服", "新功能优先体验"]
            }
        ]
    }

# ============ 原有转换接口（加用户验证） ============

# 创建临时目录
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# 清理临时文件（超过1小时）
def cleanup_old_files():
    """清理超过1小时的临时文件"""
    import time
    current_time = time.time()
    for directory in [UPLOAD_DIR, OUTPUT_DIR]:
        for file in directory.iterdir():
            if current_time - file.stat().st_mtime > 3600:
                try:
                    file.unlink()
                except:
                    pass

cleanup_old_files()


@app.get("/")
async def root():
    """API根路径"""
    return {
        "message": "文档转换API服务",
        "version": "1.1.0",
        "endpoints": {
            "/api/convert/pdf-to-word": "PDF转Word",
            "/api/convert/word-to-pdf": "Word转PDF",
            "/api/convert/image-to-pdf": "图片转PDF",
            "/api/convert/pdf-to-image": "PDF转图片",
            "/api/user/register": "用户注册",
            "/api/user/login": "用户登录",
            "/api/vip/plans": "VIP套餐"
        }
    }


@app.get("/admin")
async def admin_page():
    """管理后台页面"""
    admin_html = Path("../admin.html")
    if admin_html.exists():
        return HTMLResponse(content=admin_html.read_text(encoding="utf-8"))
    return HTMLResponse(content="<h1>Admin page not found</h1>", status_code=404)


@app.post("/api/convert/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not PDF2DOCX_AVAILABLE:
        raise HTTPException(status_code=501, detail="PDF转Word功能暂不可用")
    try:
        file_id = str(uuid.uuid4())
        pdf_path = UPLOAD_DIR / f"{file_id}.pdf"
        docx_path = OUTPUT_DIR / f"{file_id}.docx"
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        cv = Converter(str(pdf_path))
        cv.convert(str(docx_path))
        cv.close()
        pdf_path.unlink()
        return FileResponse(
            path=str(docx_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=file.filename.replace(".pdf", ".docx")
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


@app.post("/api/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="请上传.docx格式的Word文件")
        file_id = str(uuid.uuid4())
        docx_path = UPLOAD_DIR / f"{file_id}.docx"
        pdf_path = OUTPUT_DIR / f"{file_id}.pdf"
        with open(docx_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        doc = Document(str(docx_path))
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        width, height = letter
        margin = 50
        line_height = 14
        y_position = height - margin
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            title = doc.paragraphs[0].text.strip()
            c.setFont("Helvetica-Bold", 16)
            c.drawString(margin, y_position, title)
            y_position -= line_height * 2
        else:
            c.setFont("Helvetica-Bold", 16)
            c.drawString(margin, y_position, "Document")
            y_position -= line_height * 2
        c.setFont("Helvetica", 11)
        for para in doc.paragraphs[1:]:
            text = para.text.strip()
            if not text:
                y_position -= line_height
                continue
            words = text.split()
            current_line = ""
            for word in words:
                test_line = current_line + word + " "
                if c.stringWidth(test_line, "Helvetica", 11) > (width - 2 * margin):
                    if y_position < margin:
                        c.showPage()
                        y_position = height - margin
                    c.drawString(margin, y_position, current_line.strip())
                    y_position -= line_height
                    current_line = word + " "
                else:
                    current_line = test_line
            if current_line.strip():
                if y_position < margin:
                    c.showPage()
                    y_position = height - margin
                c.drawString(margin, y_position, current_line.strip())
                y_position -= line_height
        c.save()
        try:
            docx_path.unlink()
        except Exception as e:
            print(f"警告：无法删除临时Word文件: {e}")
        return FileResponse(
            path=str(pdf_path),
            media_type="application/pdf",
            filename=file.filename.replace(".docx", ".pdf")
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


@app.post("/api/convert/image-to-pdf")
async def image_to_pdf(file: UploadFile = File(...)):
    try:
        file_id = str(uuid.uuid4())
        image_path = UPLOAD_DIR / file.filename
        pdf_path = OUTPUT_DIR / f"{file_id}.pdf"
        with open(image_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        img = Image.open(image_path)
        img_width, img_height = img.size
        c = canvas.Canvas(str(pdf_path), pagesize=(img_width, img_height))
        c.drawImage(str(image_path), 0, 0, width=img_width, height=img_height)
        c.save()
        img.close()
        import time
        time.sleep(0.1)
        try:
            image_path.unlink()
        except Exception as e:
            print(f"警告：无法删除临时图片文件: {e}")
        return FileResponse(
            path=str(pdf_path),
            media_type="application/pdf",
            filename=file.filename.replace(Path(file.filename).suffix, ".pdf")
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


@app.post("/api/convert/pdf-to-image")
async def pdf_to_image(file: UploadFile = File(...)):
    if not PDF2IMAGE_AVAILABLE:
        raise HTTPException(status_code=501, detail="PDF转图片功能暂不可用")
    try:
        import zipfile
        file_id = str(uuid.uuid4())
        pdf_path = UPLOAD_DIR / f"{file_id}.pdf"
        image_dir = OUTPUT_DIR / f"{file_id}"
        image_dir.mkdir(exist_ok=True)
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        images = convert_from_path(str(pdf_path))
        # 如果只有1页，直接返回图片
        if len(images) == 1:
            image_path = image_dir / "page_1.png"
            images[0].save(str(image_path), "PNG")
            pdf_path.unlink()
            return FileResponse(
                path=str(image_path),
                media_type="image/png",
                filename=file.filename.replace(".pdf", ".png")
            )
        # 多页打包为ZIP
        zip_path = OUTPUT_DIR / f"{file_id}.zip"
        with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zf:
            for i, img in enumerate(images):
                img_bytes_io = __import__('io').BytesIO()
                img.save(img_bytes_io, 'PNG', optimize=True)
                img_bytes_io.seek(0)
                zf.writestr(f"page_{i+1}.png", img_bytes_io.read())
        pdf_path.unlink()
        # 清理图片目录
        import shutil as _shutil
        _shutil.rmtree(image_dir, ignore_errors=True)
        return FileResponse(
            path=str(zip_path),
            media_type="application/zip",
            filename=file.filename.replace(".pdf", "_images.zip")
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "service": "document-converter", "version": "1.1.0"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
